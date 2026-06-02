"""Build a single figures.docx containing Figs 1-6 with their legends."""
from __future__ import annotations
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from src import config

FIGURES = [
    {"n": 1, "title": "Figure 1. NCRB-vs-GBD suicide-mortality ratio by Indian state, 2023.",
     "image": config.OUT_FIGS / "fig1_ncrb_gbd_ratio_colour.png",
     "legend": "Choropleth (or bar chart if state geometry mismatched) of the NCRB suicide count divided by the GBD-modelled state self-harm deaths for 2023. Values below 1.0 indicate NCRB under-counts relative to GBD; values above 1.0 indicate NCRB over-counts. National aggregate ratio was 0.854."},
    {"n": 2, "title": "Figure 2. Avertable suicide burden under Scenario C (WHO global 9.0/100k), India 2023.",
     "image": config.OUT_FIGS / "fig2_avertable_scenario_C_colour.png",
     "legend": "Horizontal bar chart of avertable suicide deaths per 100,000 population if every state matched the WHO global benchmark of 9.0 suicide deaths per 100,000. Sorted ascending. States already below the WHO benchmark contribute zero."},
    {"n": 3, "title": "Figure 3. Avertable suicide deaths by Indian state under three counterfactual scenarios, 2023.",
     "image": config.OUT_FIGS / "fig3_forest_three_scenarios.png",
     "legend": "Forest-style triple-bar plot of avertable deaths under Scenarios A (Indian top-quartile, 10.04/100k), B (national average, 14.43/100k) and C (WHO global, 9.0/100k). States sorted by Scenario C magnitude."},
    {"n": 4, "title": "Figure 4. Concentration curve of Indian state-level suicide mortality, 2023.",
     "image": config.OUT_FIGS / "fig4_concentration_curve.png",
     "legend": "Concentration curve based on GBD-modelled state rates ranked low to high. Diagonal indicates perfect equality; curve below the diagonal indicates concentration of suicide mortality in higher-rate states. Erreygers Concentration Index +0.18."},
    {"n": 5, "title": "Figure 5. Joint Bayesian posterior true state suicide rates (theta) with 95% credible intervals.",
     "image": config.OUT_FIGS / "fig5_joint_bayes_state_rates.png",
     "legend": "Joint Bayesian measurement-error model (NCRB Poisson + GBD Normal observe shared theta with state-specific NCRB under-registration coefficient c) posterior mean theta per 100,000 with 95% credible intervals (black circles + bars). NCRB observed (grey squares) and GBD modelled (white triangles) shown for context. States sorted ascending by posterior mean theta. MCMC: 4 chains x 2,000 draws after 2,000 warm-up. Max R-hat 1.001. The c_s posterior captures NCRB administrative coverage state-by-state (Manipur 0.08 to Kerala 0.99); full posteriors in Supplement Table S13."},
    {"n": 9, "title": "Figure 9. SES gradient of state suicide rate, India 2023.",
     "image": config.OUT_FIGS / "fig9_ses_gradient.png",
     "legend": "State-level mean of district-level NFHS-5 (2019-21) PCA-derived SES composite (x-axis) vs NCRB 2023 suicide rate (y-axis). Spearman rho = +0.47 (p=0.012); state SII = +18.74/100k; state RII = 4.47."},
    {"n": 10, "title": "Figure 10. State x sex joint Bayesian posterior suicide rates.",
     "image": config.OUT_FIGS / "fig10_state_sex_bayes.png",
     "legend": "Posterior mean state suicide rates (per 100,000) with 95% credible intervals. Squares = male; circles = female. Global sex multiplier model (M:F = 1.46 across all 31 states). Max R-hat = 1.0001; 4 chains x 2,000 draws after 2,000 warm-up. 'Other Union Territories' aggregate excluded due to GBD's grouping."},
    {"n": 6, "title": "Figure 6. National (NCRB 2018-2024 + GBD 2020-2023) and state-level suicide-mortality trend, India.",
     "image": config.OUT_FIGS / "fig6_trend_2018_2024.png",
     "legend": "Panel (a) overlays NCRB official annual totals 2018-2024 (blue circles) with GBD modelled totals 2020-2023 (red squares) plus GBD 95% uncertainty interval; COVID-19 period shaded. NCRB rose 26.9% from 134,516 (2018) to 171,418 (2023), then plateaued at 170,746 (2024); GBD peaked at 205,876 (2021) and plateaued at 200,060 (2023). Panel (b) shows GBD-modelled trends for the top-five burden states. The national NCRB/GBD ratio rose from 0.82 (2020) to 0.86 (2023), indicating progressive NCRB administrative coverage gain."},
]


def _add_figure(doc: Document, fig: dict) -> None:
    p = doc.add_paragraph(); r = p.add_run(fig["title"]); r.bold = True; r.font.size = Pt(12)
    if fig["image"].exists():
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(fig["image"]), width=Inches(6.0))
    else:
        doc.add_paragraph().add_run(f"[Figure file not found at {fig['image']}; placeholder]").italic = True
    leg = doc.add_paragraph(); leg.add_run(fig["legend"]); leg.paragraph_format.space_after = Pt(12)
    doc.add_page_break()


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Times New Roman"; style.font.size = Pt(11)
    t = doc.add_paragraph()
    r = t.add_run("Figures with legends -- State-level reconciliation and equity decomposition of suicide mortality in India: triangulating NCRB and GBD with a Bayesian hierarchical model, 2023")
    r.bold = True; r.font.size = Pt(13)
    doc.add_paragraph("Sole author: Dr Siddalingaiah H S. Six main-text figures with embedded images and full captions, one per page; suitable for IJMR companion submission or standalone JOIAC upload.")
    doc.add_page_break()
    for fig in FIGURES:
        _add_figure(doc, fig)
    out_path = config.MANUSCRIPT / "figures.docx"
    doc.save(out_path)
    print(f"[figures_docx] wrote {out_path}")


if __name__ == "__main__":
    main()
